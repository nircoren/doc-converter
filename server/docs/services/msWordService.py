import copy
import sys
sys.path.append("../result-tests")

from docx import Document
from docx.shared import Cm,Inches
from docx.oxml.shared import OxmlElement 
from docx.oxml.ns import qn,nsdecls
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.text.run import CT_R
from PIL import Image
from docx.oxml import parse_xml
from docx.shared import Pt
import uuid

def tif_to_png(tif_file_path):
    with Image.open(tif_file_path) as img:
        # Check if the TIF file contains multiple pages (frames)
        if img.n_frames > 1:
            for frame in range(img.n_frames):
                img.seek(frame)
                img.save(f"page_{frame + 1}.png")
        else:
            # If the TIF file has only one page
            img.save("page_1.png")


def create_new_doc(url):
    return Document(url)

def save_doc(doc,font_style):
    doc_id = str(uuid.uuid4())
    url = f'./results/res-{doc_id}.docx'
    font = doc.styles['Normal'].font
    font.name = font_style['name']
    font.size = Pt(font_style['size'])
    doc.save(url)
    return doc_id

def insert_text_to_table(doc,insertion,text):
    # doc = Document(targetUrl)
    table = doc.tables[insertion['indexes']['table_index']]
    target_row_index = insertion['indexes']['row']
    target_column_index = insertion['indexes']['column']
    target_cell = table.cell(target_row_index, target_column_index)
    target_cell.paragraphs[0].text = text

def __set_cell_style(paragraph,insertion,text):
    if 'style' in insertion and 'font_size' in insertion['style']:
        paragraph.style.font.size = Pt(insertion['style']['font_size'])

# insert text in table. copy style from first cell to second cell
def copy_text_style(source_run, target_run):
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb
    target_run.font.highlight_color = source_run.font.highlight_color
    target_run.font.bold = source_run.font.bold
    target_run.font.italic = source_run.font.italic
    target_run.font.underline = source_run.font.underline
    target_run.font.strike = source_run.font.strike


def copy_cell_text_style(source_cell, target_cell):
    for source_paragraph in source_cell.paragraphs:
        target_paragraph = target_cell.add_paragraph()
        for source_run in source_paragraph.runs:
            target_run = target_paragraph.add_run()
            copy_text_style(source_run, target_run)

def main():
    doc = Document('test-form.docx')
    first_table = doc.tables[0]
    cell = first_table.rows[0].cells[1]
    copy_cell_text_style(first_table.rows[0].cells[0], cell)
    doc.save('test-form1534534523132133.docx')  


def insert_word_after_sentence(doc_path, target_sentence, word_to_insert):
    doc = Document(doc_path)
    found = False

    for paragraph in doc.paragraphs:
        if target_sentence in paragraph.text:
            paragraph.text = paragraph.text.replace(target_sentence, f"{target_sentence} {word_to_insert}")
            found = True
            break

    if not found:
        print("Target sentence not found in the document.")

def check_box(doc,insertion):
    paragraph_includes = insertion['paragraph_includes']
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if paragraph_includes in paragraph.text:
            target_paragraph = paragraph
            break

    # check checkbox
    if target_paragraph:
        x = paragraph._element.xpath('.//w:checkBox')
        x[insertion['index']].append(checkedElement())

    for elem in paragraph._element.xpath('.//w:checkBox'):
        checked = parse_xml(r'<w:checked {}/>'.format(nsdecls('w')))
        elem.getparent().insert(0, checked)
        break
#         x[index].append(checkedElement())
    

def insert_text_to_paragraph(doc,insertion,text):
    text_before = insertion['text_before']
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if text_before in paragraph.text:
            target_paragraph = paragraph
            break

    if target_paragraph:
        under_line_symbol = '\u2002'
        full_text = ''
        text_inserted = False
        for run in target_paragraph.runs:
            full_text += run.text
            if text_before in full_text and not text_inserted:
                run.text += ' '
                add_new_run_to_paragraph(run,target_paragraph,insertion, text)
                text_inserted = True
            if run.text == under_line_symbol:
                run.text = ''
        # else: 
            # target_paragraph.runs[0].text = target_paragraph.text[:idx] + ' ' + text_to_insert['text'] + '' + target_paragraph.text[idx:]


def add_new_run_to_paragraph(run,target_paragraph,insertion, text):
    new_run_element = target_paragraph._element._new_r()
    run._element.addnext(new_run_element)
    new_run = Run(new_run_element, run._parent)
    # ---do things with new_run, e.g.---
    new_run.text = text
    copy_run_style(run, new_run)
    add_run_style(insertion, new_run,text)
    
def insert_word_in_run(inline, i, start_run_text, mid_run_text,end_run_text):
    start_run,mid_run,end_run =  [copy.deepcopy(inline[i]) for _ in range(3)]
    start_run.text,mid_run.text,end_run.text 
    
def add_run_style(insertion, run,text):
    if('style' not in insertion):
        return
    for key in insertion['style']:
        setattr(run, key, insertion['style'][key])
    # for key in target_font_style:
    #     setattr(run.font, key, target_font_style[key])

def copy_run_style(source_run, target_run):
    # Copy font properties
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    # target_run.font.bold = source_run.font.bold
    # target_run.font.italic = source_run.font.italic
    # target_run.font.underline = source_run.font.underline
    # target_run.font.strike = source_run.font.strike
    # target_run.font.superscript = source_run.font.superscript
    # target_run.font.subscript = source_run.font.subscript
    # target_run.font.color.rgb = source_run.font.color.rgb
    # target_run.font.highlight_color = source_run.font.highlight_color

def replace_text(text, to_replace):
    if to_replace == 'line':
        return text.replace("\n", "")
# Example usage
# document_path = "test-form.docx"
# target_sentence = "אני מעיד כי בתאריך"
# word_to_insert_after_sentence = "10.01.2021"

# insert_word_after_sentence(document_path, target_sentence, word_to_insert_after_sentence)


#  check box tests:
# def yesNoCheck(yes_no,tableIdx,coords):
#     if yes_no == 'y':
#         index = 0
#         x = doc.tables[tableIdx].cell(coords[0],coords[1])._element.xpath('.//w:checkBox')
#         x[index].append(checkedElement())
#     elif yes_no == 'n':
#         index = 1
#         x = doc.tables[tableIdx].cell(coords[0],coords[1])._element.xpath('.//w:checkBox')
#         x[index].append(checkedElement())
#     else:
#         print ("value was neither yes or no")
#         pass

def checkedElement():
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'),"true")
    return elm